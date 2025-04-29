import pandas as pd
import numpy as np
import matplotlib.pyplot as plt
from docx import Document
from docx.shared import Inches
import io


class DeviceAnalyzer:
    """
    Класс для определения состояния устройств:
    - off: потребление == 0
    - active: потребление >= q_active квантиля среди ненулевых
    - idle: все остальные значения >0 и < активного порога

    Возвращает проценты и стататику максимальных непрерывных участков.
    """
    def __init__(self, data: pd.DataFrame, q_active: float = 0.75):
        self.data = data
        self.q_active = q_active

    def _compute_mask_and_stats(self, series_full: pd.Series) -> dict:
        # Сглаживание медианным фильтром по полным данным
        smooth = series_full.rolling(window=1, center=True, min_periods=1).median()
        # Определяем маски
        mask_off = series_full == 0
        nonzero = series_full[series_full > 0]
        total = series_full.notna().sum()
        if total == 0:
            return {'off_percent': 0.0,
                    'idle_percent': 0.0,
                    'active_percent': 0.0,
                    'max_off_run': 0,
                    'max_idle_run': 0,
                    'max_active_run': 0}
        # Если нет ненулевых — всё off
        if nonzero.empty:
            return {'off_percent': mask_off.sum() / total * 100,
                    'idle_percent': 0.0,
                    'active_percent': 0.0,
                    'max_off_run': self._max_run(mask_off),
                    'max_idle_run': 0,
                    'max_active_run': 0}
        # Определяем порог активности
        th_active = nonzero.quantile(self.q_active)
        mask_active = smooth >= th_active
        mask_idle = (~mask_off) & (~mask_active) & series_full.notna()
        # Вычисляем проценты
        off_p = mask_off.sum() / total * 100
        active_p = mask_active.sum() / total * 100
        idle_p = mask_idle.sum() / total * 100
        # Статистика непрерывных участков
        return {'off_percent': off_p,
                'idle_percent': idle_p,
                'active_percent': active_p,
                'max_off_run': self._max_run(mask_off),
                'max_idle_run': self._max_run(mask_idle),
                'max_active_run': self._max_run(mask_active)}

    @staticmethod
    def _max_run(mask: pd.Series) -> int:
        runs = mask.astype(int).groupby((mask != mask.shift()).cumsum()).sum()
        return int(runs.max()) if not runs.empty else 0

    def get_states(self) -> dict:
        states = {}
        for dev in self.data.columns:
            series_full = self.data[dev]
            stats = self._compute_mask_and_stats(series_full)
            states[dev] = stats
        return states


class EnergyReportGenerator:
    """
    Класс для генерации отчёта по потреблению электроэнергии из Excel в Word
    """
    def __init__(self, filepath: str, sheet_name: str):
        self.filepath = filepath
        self.sheet_name = sheet_name
        self.raw = None
        self.data = None
        self.daily = None
        self.doc = Document()

    def load_and_prepare(self):
        # Загрузка и предобработка данных
        self.raw = pd.read_excel(self.filepath, sheet_name=self.sheet_name, skiprows=1)
        self.raw.columns = self.raw.columns.astype(str).str.strip()
        self.raw['DateTime'] = pd.to_datetime(
            self.raw['Дата'].astype(str) + ' ' + self.raw['Время'].astype(str), dayfirst=True
        )
        self.raw.set_index('DateTime', inplace=True)
        # Оставляем только числовые данные
        self.data = self.raw.drop(columns=['Дата', 'Время']).apply(pd.to_numeric, errors='coerce')

    def generate_daily(self):
        # Суммарное потребление по дням
        self.daily = self.data.resample('D').sum()

    def add_heading(self):
        start = self.daily.index.min().strftime('%d.%m.%Y')
        end = self.daily.index.max().strftime('%d.%m.%Y')
        self.doc.add_heading(
            f'Отчет о потреблении электроэнергии за период с {start} по {end}', level=0
        )
        self.doc.add_paragraph(
            'В отчете представлены суммарные и детальные показатели потребления электричества оборудованием.'
        )

    def plot_all_devices(self):
        plt.figure(figsize=(14, 8))
        for col in self.daily.columns:
            plt.plot(self.daily.index, self.daily[col], label=col)
        plt.title('Суточное потребление: все устройства')
        plt.xlabel('Дата')
        plt.ylabel('кВт·ч')
        plt.grid(True)
        plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left')
        plt.tight_layout()
        buf = io.BytesIO()
        plt.savefig(buf, format='png')
        buf.seek(0)
        plt.close()

        self.doc.add_picture(buf, width=Inches(6))
        self.doc.add_paragraph('Рисунок 1. Суточное потребление всех устройств.')

    def plot_top10_and_table(self):
        total = self.daily.sum().sort_values(ascending=False)
        top10 = total.head(10).index

        # График топ-10
        plt.figure(figsize=(12, 5))
        for col in top10:
            plt.plot(self.daily.index, self.daily[col], label=col)
        plt.title('Суточное потребление: топ-10 устройств')
        plt.xlabel('Дата')
        plt.ylabel('кВт·ч')
        plt.grid(True)
        plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left')
        plt.tight_layout()
        buf = io.BytesIO(); plt.savefig(buf, format='png'); buf.seek(0); plt.close()
        self.doc.add_picture(buf, width=Inches(6))
        self.doc.add_paragraph('Рисунок 2. Топ-10 потребителей электроэнергии.')

        # Таблица топ-10
        self.doc.add_paragraph('Таблица 1. Суммарное потребление по топ-10 устройствам:')
        table = self.doc.add_table(rows=1, cols=2, style='Table Grid')
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text = 'Устройство', 'Потребление (кВт·ч)'
        for name, val in total.head(10).items():
            row = table.add_row().cells
            row[0].text = name
            row[1].text = f"{val:.2f}"

    def plot_categories(self):
        # Автоматическая классификация по названиям
        def classify(cols):
            cats = {'PzS_12V': [], 'China': [], 'SM': [], 'MO': [], 'BG': [], 'DIG': [], 'CP-300': [], 'Other': []}
            for c in cols:
                c_low = c.lower()
                if 'pzs' in c_low and '12v' in c_low:
                    cats['PzS_12V'].append(c)
                elif 'china' in c_low:
                    cats['China'].append(c)
                elif ' sm' in c_low or 'sm ' in c_low:
                    cats['SM'].append(c)
                elif ' mo' in c_low or 'mo ' in c_low:
                    cats['MO'].append(c)
                elif ' bg' in c_low or 'bg ' in c_low:
                    cats['BG'].append(c)
                elif 'dig' in c_low:
                    cats['DIG'].append(c)
                elif 'cp-300' in c_low:
                    cats['CP-300'].append(c)
                else:
                    cats['Other'].append(c)
            return {k: v for k, v in cats.items() if v}

        cats = classify(self.daily.columns)
        cat_df = pd.DataFrame({cat: self.daily[cols].sum(axis=1) for cat, cols in cats.items()})

        plt.figure(figsize=(12, 6))
        for col in cat_df.columns:
            plt.plot(cat_df.index, cat_df[col], label=col)
        plt.title('Потребление по категориям оборудования')
        plt.xlabel('Дата')
        plt.ylabel('кВт·ч')
        plt.grid(True)
        plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left')
        plt.tight_layout()
        buf = io.BytesIO(); plt.savefig(buf, format='png'); buf.seek(0); plt.close()

        self.doc.add_picture(buf, width=Inches(6))
        self.doc.add_paragraph('Рисунок 3. Потребление по категориям оборудования.')

    def temporal_analysis(self):
        # Среднее по часам суток для топ-10
        hourly = self.data.resample('H').mean()
        total = self.daily.sum().sort_values(ascending=False)
        top10 = total.head(10).index
        typical = hourly.groupby(hourly.index.hour)[top10].mean()

        plt.figure(figsize=(14, 7))
        for col in typical.columns:
            plt.plot(typical.index, typical[col], label=col)
        plt.title('Среднее потребление по часам суток')
        plt.xlabel('Час дня')
        plt.ylabel('кВт·ч')
        plt.xticks(np.arange(0, 24, 1))
        plt.grid(True)
        plt.legend(bbox_to_anchor=(1.05, 1), loc='upper left')
        plt.tight_layout()
        buf = io.BytesIO(); plt.savefig(buf, format='png'); buf.seek(0); plt.close()

        self.doc.add_picture(buf, width=Inches(6))
        self.doc.add_paragraph('Рисунок 4. Среднее потребление по часам суток (топ-10).')

    def anomaly_analysis(self):
        # Анализ аномалий потребления
        sigma, window, top_n = 2, 24, 10
        self.doc.add_heading('Аномалии потребления', level=1)
        def analyze(name, series):
            rm = series.rolling(window).mean()
            rs = series.rolling(window).std()
            mask = (series > rm + sigma*rs) | (series < rm - sigma*rs)
            anomalies = series[mask]
            if anomalies.empty:
                return None
            dev = np.abs(anomalies - rm[anomalies.index])
            return {'count': len(anomalies), 'max_dev': dev.max(), 'mean_dev': dev.mean(), 'total_dev': dev.sum()}

        results = []
        for dev in self.data.columns:
            res = analyze(dev, self.data[dev].dropna())
            if res:
                results.append((dev, res))
        results.sort(key=lambda x: x[1]['total_dev'], reverse=True)
        top = results[:top_n]

        table = self.doc.add_table(rows=1, cols=5, style='Table Grid')
        hdr = table.rows[0].cells
        for i, txt in enumerate(['Устройство', 'Кол-во аномалий', 'Макс откл.', 'Сред откл.', 'Сумма откл.']):
            hdr[i].text = txt
        for dev, r in top:
            row = table.add_row().cells
            row[0].text = dev
            row[1].text = str(r['count'])
            row[2].text = f"{r['max_dev']:.2f}"; row[3].text = f"{r['mean_dev']:.2f}"; row[4].text = f"{r['total_dev']:.2f}"

    def device_state_analysis(self):
        # Вставка анализа состояния оборудования
        self.doc.add_heading('Состояние оборудования', level=1)
        analyzer = DeviceAnalyzer(self.data)
        states = analyzer.get_states()
        table = self.doc.add_table(rows=1, cols=4, style='Table Grid')
        hdr = table.rows[0].cells
        hdr[0].text, hdr[1].text, hdr[2].text, hdr[3].text = (
            'Устройство', 'Отключено (%)', 'Простаивает (%)', 'Работает (%)'
        )
        for dev, m in states.items():
            row = table.add_row().cells
            row[0].text = dev
            row[1].text = f"{m['off_percent']:.1f}"
            row[2].text = f"{m['idle_percent']:.1f}"
            row[3].text = f"{m['active_percent']:.1f}"

    def save(self, out: str = 'report.docx'):
        self.doc.save(out)

    def run(self):
        self.load_and_prepare()
        self.generate_daily()
        self.add_heading()
        self.plot_all_devices()
        self.plot_top10_and_table()
        self.plot_categories()
        self.temporal_analysis()
        self.anomaly_analysis()
        self.device_state_analysis()
        self.save()


if __name__ == '__main__':
    generator = EnergyReportGenerator(
        filepath="Задание_2_Потребление_электроэнергии_Апрель.xlsx",
        sheet_name="2025-04-01-00-00-00-e"
    )
    generator.run()
